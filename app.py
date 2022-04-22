import smtplib
import email, smtplib, ssl
import docx
import pandas as pd
import requests
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
from flask import Flask, render_template, url_for, request
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import requests
import json
import pandas as pd
import numpy as np
from gql import Client, gql
from gql.transport.requests import RequestsHTTPTransport
import re
import yfinance as yf
import datetime
from datetime import datetime, date,timedelta
from timeit import default_timer as timer


app = Flask(__name__)


pd.set_option('display.max_columns', 50)
pd.set_option('display.max_rows', 50)
pd.options.mode.chained_assignment = None


uri = "https://demo.api.alpha-sense.com/gql"
headers = {
    'x-api-key': '0OcnshpPJ36zl89j4IXFynTzMKu0j3zanGftgqI4',
     
    'clientid': 'as-api-demo',
    'Authorization': 'Bearer 7d7ff62e-56df-430b-9c0b-ff1c8c678b8e',
    'Content-Type': 'application/json'
}
transport = RequestsHTTPTransport(uri, headers)
client = Client(transport=transport)

df = pd.read_excel(pd.read_excel("https://docs.google.com/spreadsheets/d/e/2PACX-1vQJ2q5Ntw4qfoMoUG1u9M_eDUyyhwFzj2xFUGIXsdA31x_sprT3o_c903hKYRkXbA/pub?output=xlsx",'keywords')
df=df.fillna(' ')

df_coef = pd.read_excel('https://docs.google.com/spreadsheets/d/e/2PACX-1vQJ2q5Ntw4qfoMoUG1u9M_eDUyyhwFzj2xFUGIXsdA31x_sprT3o_c903hKYRkXbA/pub?output=xlsx','coefs')

df_dates = pd.read_excel('https://docs.google.com/spreadsheets/d/e/2PACX-1vQJ2q5Ntw4qfoMoUG1u9M_eDUyyhwFzj2xFUGIXsdA31x_sprT3o_c903hKYRkXbA/pub?output=xlsx','dates')

today = date.today()
fromdate = today-timedelta(days=90)
new_date = f'"{today.strftime("%Y-%m-%d")}"'
from_date= f'"{fromdate.strftime("%Y-%m-%d")}"'
last_date=df_dates['dates'][1] #now, new date becomes last date

#print(from_date,last_date,new_date)

#df_dates2=pd.DataFrame({'description':["from_date","last_date","new_date"],'dates':[from_date,last_date,new_date]})

#writer = pd.ExcelWriter(r'test.xlsx')
#df.to_excel(writer, sheet_name = 'keywords')
#df_coef.to_excel(writer, sheet_name = 'coefs')
#df_dates2.to_excel(writer, sheet_name = 'dates')

#writer.save()

#for now we will not use this cell and we will overwrite the dates

from_date="2021-09-15"
last_date="2022-02-15"
new_date="2022-03-15"

ind_name_list=df['Industry Name'].tolist()
ind_code_list=df["Industry code"].tolist()
pm_list=df['Product Margin'].tolist()
pm_list.append("Product Margin")
oc_list=df['Operational Complexity'].tolist()
s_list=df['Sustainability'].tolist()
d_list=df['Design'].tolist()
co_list=df['Customer Obsession'].tolist()
lp_list=df['Launch Products'].tolist()
o_list=df['Other'].tolist()
problems_list=list(df.columns)[-7:]

ind_list=list()
for (idx,val) in enumerate(ind_name_list):
    ind_list+=[[ind_code_list[idx], ind_name_list[idx],pm_list[idx],oc_list[idx],s_list[idx],d_list[idx],co_list[idx],lp_list[idx],o_list[idx]]]


re.split(";",str(ind_list[0][0]))

for i in ind_list:
    i[0]=re.split(";",str(i[0]))
    i[2]=re.split(";",i[2])
    i[3]=re.split(";",i[3])
    i[4]=re.split(";",i[4])
    i[5]=re.split(";",i[5])
    i[6]=re.split(";",i[6])
    i[7]=re.split(";",i[7])
    i[8]=re.split(";",i[8])

all_search_requests=[]
for i in range(len(ind_list)):
    for j in range(len(ind_list[i][0])):
        ind_code=ind_list[i][0][j]
        ind_name=ind_list[i][1]
        for k in range(len(ind_list[i])-2):
                keywords= ind_list[i][k+2]
                problem=problems_list[k]
                temp=[ind_code,keywords,problem,ind_name]
                all_search_requests.append(temp)


n=100
n_max=len(all_search_requests)
if n>n_max:
    n=n_max
iterator=all_search_requests[:n]


date_list=[last_date,new_date]
comparison={}
result_list=[]

for i in iterator:

    start = timer()

    res=[]
    r = []
    results = []
    #initialization
    res=list()

    var = 100
    flag=0
    var_iter=0

    num_per_call = 100

    cursor = ""

    # Open output text file
    f = open("output.json", 'w')

    industry =[]
    industry.append(str(i[0])) #get the industry code
    keywords= i[1] # get the keywords list


    querytext=''
    for j in keywords:
        if len(re.split(" ",j))>1:
            querytext= querytext + 'NEAR0('+j+') OR '        
        else:
            querytext= querytext +j+ ' OR '

    if querytext[-3:]=='OR ':
        querytext=querytext[:-4]

    querytext = f"({querytext})" + ' negative'
    #print(i[0]+'\t'+querytext)

df = pd.read_excel('https://docs.google.com/spreadsheets/d/e/2PACX-1vQJ2q5Ntw4qfoMoUG1u9M_eDUyyhwFzj2xFUGIXsdA31x_sprT3o_c903hKYRkXbA/pub?output=xlsx')
df=df.fillna(' ')
ind_name_list=df['Industry Name'].tolist()
ind_code_list=df["Industry code"].tolist()
pm_list=df['Product Margin'].tolist()
pm_list.append("Product Margin")
oc_list=df['Operational Complexity'].tolist()
s_list=df['Sustainability'].tolist()
d_list=df['Design'].tolist()
co_list=df['Customer Obsession'].tolist()
lp_list=df['Launch Products'].tolist()
o_list=df['Other'].tolist()
problems_list=list(df.columns)[-7:]

ind_list=list()
for (idx,val) in enumerate(ind_name_list):
    ind_list+=[[ind_code_list[idx], ind_name_list[idx],pm_list[idx],oc_list[idx],s_list[idx],d_list[idx],co_list[idx],lp_list[idx],o_list[idx]]]


for i in ind_list:
    i[0]=re.split(";",str(i[0]))
    i[2]=re.split(";",i[2])
    i[3]=re.split(";",i[3])
    i[4]=re.split(";",i[4])
    i[5]=re.split(";",i[5])
    i[6]=re.split(";",i[6])
    i[7]=re.split(";",i[7])
    i[8]=re.split(";",i[8])

all_search_requests=[]
for i in range(len(ind_list)):
    for j in range(len(ind_list[i][0])):
        ind_code=ind_list[i][0][j]
        ind_name=ind_list[i][1]
        for k in range(len(ind_list[i])-2):
                keywords= ind_list[i][k+2]
                problem=problems_list[k]
                temp=[ind_code,keywords,problem,ind_name]
                all_search_requests.append(temp)

date_list=[last_date,new_date]
comparison={}
result_list=[]

len_iterator=len(iterator)
counter=0
counter_int=0

for i in iterator:
    #print('Connecting to Alphasense')
    for datei in date_list:

        start = timer()

        res=[]
        r = []
        results = []
        #initialization
        res=list()

        var = 100
        flag=0
        var_iter=0

        num_per_call = 100

        cursor = ""

        # Open output text file
        f = open("output.json", 'w')

        industry =[]
        industry.append(str(i[0])) #get the industry code
        keywords= i[1] # get the keywords list


        querytext=''
        for j in keywords:
            if len(re.split(" ",j))>1:
                querytext= querytext + 'NEAR0('+j+') OR '        
            else:
                querytext= querytext +j+ ' OR '

        if querytext[-3:]=='OR ':
            querytext=querytext[:-4]

        querytext = f"({querytext})" + ' negative'


        while (flag!=1) and (var_iter<var):
            # Change these parameters to edit the inputs
            params = {
                    "filter":{
                            "keyword":{
                                    "query": querytext
                            },
                            "date":{
                                    "customRange": {
                                            "from":from_date,
                                            "to":datei,
                                    }
                            },
                            "industries": industry,
                            "countries": ["US"]
                    },
                    "sorting": {
                            "field":"DATE",
                            "direction":"DESC"
                    },
                    "limit": num_per_call,
                    "cursor": cursor
            }

            query = gql('''
            query alphasensequery ($filter: SearchFilter!, $limit: Int!, $sorting: SearchSorting!, $cursor: String) { 
                    search (filter: $filter, limit: $limit, sorting: $sorting, cursor: $cursor){                
                            cursor
                            totalCount         
                            documents {
                                    companies {
                                            name
                                            primaryTickerCode
                                            isin
                                    }
                            }
                    }
            }''')

            result = client.execute(query, variable_values=params)

            cursor = result["search"]["cursor"]

            if cursor==None:
                flag=1
            var_iter+=1
            n_res=result["search"]["totalCount"]
            res+=result["search"]["documents"]
            comparison[datei]=res        
    if comparison[date_list[0]]==comparison[date_list[1]]:
        continue

    else:
        comps={}
        tickers={}

        last_res=comparison[date_list[1]]
        for k in range(len(last_res)):
            comp_name=last_res[k]['companies'][0]['name']
            comp_id=last_res[k]['companies'][0]['primaryTickerCode']
            if (comp_name,comp_id) in comps:
                comps[(comp_name,comp_id)]+=1
            else:
                comps[(comp_name,comp_id)]=1


        df=pd.DataFrame()
        df=pd.DataFrame.from_dict(comps, orient='index',columns=["N_docs"])
        df.reset_index(inplace=True)
        df[["company","ticker"]]=pd.DataFrame(df['index'].tolist(), index=df.index)
        df.dropna(inplace=True)
        df=df[['company','ticker','N_docs']]
        df['keywords']=querytext
        df.sort_values(by=['N_docs'], ascending=False)
        df1=df
        df=df.head(3).sort_values('N_docs', ascending=False)
        df=df.reset_index()
        #print(querytext)
        #print(df[["company","ticker","N_docs"]])
        temp_list=[]
        for l in range(df['N_docs'].size):
            comp_list=[]
            comp_list.append(df['company'][l])
            comp_list.append(df['ticker'][l])
            comp_list.append(i[2])
            comp_list.append(df['keywords'][l])
            comp_list.append(df['keywords'][l][1:-10])
            temp_list.append(comp_list)
        result_list.append(temp_list)

#
#
#

n=1 #for demonstration

#
#
#

from datetime import date
def diff_month(d1, d2):
        return (d1.year - d2.year) * 12 + d1.month - d2.month
def return_company_revenue(ticker):
        try:
            return yf.Ticker(ticker).info['totalRevenue']

        except:
            pass

def return_market_cap(ticker):
        try:
            return yf.Ticker(ticker).info['marketCap']
        except:
            pass

input_list=result_list    

len_list=len(input_list)
counter=4
counter_int=4

df_final=pd.DataFrame(columns=["Company", "Ticker","Problem","Negative Query","Query","Revenue","Market Capitalization","Last Month negative documents","Last Quarter negative documents","Last Year negative documents","Last Month total documents","Last Quarter total documents","Last Year total documents","Last Month negative mentions","Last Quarter negative mentions","Last Year negative mentions","Last Month total mentions","Last Quarter total mentions","Last Year total mentions","Sentiment Score","Sentiment Change","Peer Tickers","Total Score"])


leni=len(input_list)
for h in input_list:

    i=1
    #print(f'\nStep {i} / {leni}')
    i+=1
    #df_email=pd.DataFrame
    for list_sel_comp in h:


        df_score=pd.DataFrame([list_sel_comp], columns=["Company", "Ticker","Problem","Negative Query","Query"])
        df_score['Revenue'] = df_score['Ticker'].apply(lambda x: return_company_revenue(x))
        #df_score['Market Capitalization']=df_score['Ticker'].apply(lambda x: return_market_cap(x))

        df_score["Last Month negative documents"]=""
        df_score["Last Quarter negative documents"]=""
        df_score["Last Year negative documents"]=""
        df_score["Last Month total documents"]=""
        df_score["Last Quarter total documents"]=""
        df_score["Last Year total documents"]=""
        df_score["Last Month negative mentions"]=""
        df_score["Last Quarter negative mentions"]=""
        df_score["Last Year negative mentions"]=""
        df_score["Last Month total mentions"]=""
        df_score["Last Quarter total mentions"]=""
        df_score["Last Year total mentions"]=""
        df_score["Sentiment Score"]=""
        df_score["Sentiment Change"]=""
        df_score["Peer Tickers"]=""
        df_score["Total Score"]=""

        var = 100
        flag=0
        var_iter=0
        res_neg=list()
        total_count_neg=0
        total_count_neg_1m=0
        total_count_neg_3m=0

        num_per_call = 100

        cursor = ""
        n_docs=0
        nb_docs_1m=0
        nb_docs_3m=0


        f = open("output.json", 'w')

        while (flag!=1) and (var_iter < var):
            params = {
                    "filter":{
                            "keyword":{
                                    "query": list_sel_comp[3]
                            },
                            "date":{
                                    "preset": "LAST_12_MONTHS"
                            },
                            "companies": {
                                    "include": [list_sel_comp[1]]            
                            } 
                    },
                    "sorting": {
                            "field":"DATE",
                            "direction":"DESC"
                    },
                    "limit": num_per_call,
                    "cursor": cursor
            }

            query = gql('''
            query alphasensequery ($filter: SearchFilter!, $limit: Int!, $sorting: SearchSorting!, $cursor: String) { 
                    search (filter: $filter, limit: $limit, sorting: $sorting, cursor: $cursor){                
                            cursor
                            totalCount         
                            documents {
                                    id
                                    title
                                    companies {
                                            name
                                            primaryTickerCode
                                            isin
                                    }
                                    releasedAt
                                    sentiment {
                                            net
                                            change
                                            totalPositiveCount
                                            totalNegativeCount
                                            totalStatements
                                    }
                                    snippets {
                                            hitsCount
                                            statementsCount
                                    }
                            }
                    }
            }''')
            # Calls the API
            result = client.execute(query, variable_values=params)

            # Resets the cursor for pagination
            cursor = result["search"]["cursor"]
            if cursor==None:
                flag=1

            var_iter+=1

            res_neg+=result["search"]["documents"]
            if var_iter==1:
                nb_res=result["search"]["totalCount"]

        df1=df_score.copy()
        df_score["Last Year negative documents"][(df_score["Company"]==[list_sel_comp[0]]) & (df_score["Problem"]==[list_sel_comp[2]])]=nb_res
        #print('Negative Docments: '+str(df_score["Last Year negative documents"][0]))
        #To be improved
        today=date.today()
        for j in range(len(res_neg)):
            total_count_neg+=res_neg[j]["snippets"]["statementsCount"]
            last_period=res_neg[j]["releasedAt"]
            date_doc = datetime.fromtimestamp(last_period / 1e3)
            if (diff_month(today,date_doc)<=1):
                nb_docs_1m+=1
                total_count_neg_1m+=res_neg[j]["snippets"]["statementsCount"]
            if (diff_month(today,date_doc)<=3):
                nb_docs_3m+=1
                total_count_neg_3m+=res_neg[j]["snippets"]["statementsCount"]


        df_score["Last Month negative documents"][(df_score["Company"]==[list_sel_comp[0]]) & (df_score["Problem"]==[list_sel_comp[2]])]=nb_docs_1m
        df_score["Last Quarter negative documents"][(df_score["Company"]==[list_sel_comp[0]]) & (df_score["Problem"]==[list_sel_comp[2]])]=nb_docs_3m
        df_score["Last Month negative mentions"][(df_score["Company"]==[list_sel_comp[0]]) & (df_score["Problem"]==[list_sel_comp[2]])]=total_count_neg_1m
        df_score["Last Quarter negative mentions"][(df_score["Company"]==[list_sel_comp[0]]) & (df_score["Problem"]==[list_sel_comp[2]])]=total_count_neg_3m
        df_score["Last Year negative mentions"][(df_score["Company"]==[list_sel_comp[0]]) & (df_score["Problem"]==[list_sel_comp[2]])]=total_count_neg

        # Change this varaible for pagination
        var = 100
        flag=0
        var_iter=0
        res_tot=list()
        total_count_tot=0
        total_count_tot_1m=0
        total_count_tot_3m=0

        # Change this variable for the number of results per API call
        num_per_call = 100

        # Initializing the cursor for pagination **Cursor always starts as an empty string!**
        cursor = ""
        n_docs=0
        nb_docs_1m=0
        nb_docs_3m=0



        f = open("output.json", 'w')

        # Loops for pagination
        while (flag!=1) and (var_iter < var):
            # Change these parameters to edit the inputs
            params = {
                    "filter":{
                            "keyword":{
                                    "query": list_sel_comp[4]
                            },
                            "date":{
                                    "preset": "LAST_12_MONTHS"
                            },
                            "companies": {
                                    "include": [list_sel_comp[1]]             
                            } 
                    },
                    "sorting": {
                            "field":"DATE",
                            "direction":"DESC"
                    },
                    "limit": num_per_call,
                    "cursor": cursor
            }

            # Change this if you want to change the output
            query = gql('''
            query alphasensequery ($filter: SearchFilter!, $limit: Int!, $sorting: SearchSorting!, $cursor: String) { 
                    search (filter: $filter, limit: $limit, sorting: $sorting, cursor: $cursor){                
                            cursor
                            totalCount         
                            documents {
                                    id
                                    title
                                    companies {
                                            name
                                            primaryTickerCode
                                            isin
                                    }
                                    releasedAt
                                    sentiment {
                                            net
                                            change
                                            totalPositiveCount
                                            totalNegativeCount
                                            totalStatements
                                    }
                                    snippets {
                                            hitsCount
                                            statementsCount
                                    }
                            }
                    }
            }''')

            # Calls the API
            result = client.execute(query, variable_values=params)

            # Resets the cursor for pagination
            cursor = result["search"]["cursor"]
            if cursor==None:
                flag=1

            var_iter+=1

            res_tot+=result["search"]["documents"]
            if var_iter==1:
                nb_res=result["search"]["totalCount"]


        df_score["Last Year total documents"][(df_score["Company"]==[list_sel_comp[0]]) & (df_score["Problem"]==[list_sel_comp[2]])]=nb_res


        for j in range(len(res_tot)):
            total_count_tot+=res_tot[j]["snippets"]["statementsCount"]
            last_period=res_tot[j]["releasedAt"]
            date_doc = datetime.fromtimestamp(last_period / 1e3)
            if (diff_month(today,date_doc)<=1):
                nb_docs_1m+=1
                total_count_tot_1m+=res_tot[j]["snippets"]["statementsCount"]
            if (diff_month(today,date_doc)<=3):
                nb_docs_3m+=1
                total_count_tot_3m+=res_tot[j]["snippets"]["statementsCount"]

        df_score["Last Month total documents"][(df_score["Company"]==[list_sel_comp[0]]) & (df_score["Problem"]==[list_sel_comp[2]])]=nb_docs_1m
        df_score["Last Quarter total documents"][(df_score["Company"]==[list_sel_comp[0]]) & (df_score["Problem"]==[list_sel_comp[2]])]=nb_docs_3m
        df_score["Last Month total mentions"][(df_score["Company"]==[list_sel_comp[0]]) & (df_score["Problem"]==[list_sel_comp[2]])]=total_count_tot_1m
        df_score["Last Quarter total mentions"][(df_score["Company"]==[list_sel_comp[0]]) & (df_score["Problem"]==[list_sel_comp[2]])]=total_count_tot_3m
        df_score["Last Year total mentions"][(df_score["Company"]==[list_sel_comp[0]]) & (df_score["Problem"]==[list_sel_comp[2]])]=total_count_tot
        df_final=df_final.append(df_score, ignore_index=True)
        df_final=df_final.drop(columns=["Negative Query","Query","Sentiment Score","Sentiment Change"]);

a=pd.Series()
for i in df_final:
    try:
        b=df_final[f'{i}']*df_coef[f'{i}'][0]
        a = a.add(b, fill_value=0)
    except:
        pass
df_final['Total Score']=a


df_final["Monthly Score"] = ""
df_final["Quarterly Score"] = ""
df_final["Yearly Score"] = ""
df_final["Urgent Need Reqd"] = ""


try:
    df_final["Monthly Score"] = df_final["Last Month negative mentions"] / df_final["Last Month total documents"]
except:
    df_final["Monthly Score"] = 0

try:
    df_final["Quarterly Score"] = df_final["Last Quarter negative mentions"] / df_final["Last Quarter total documents"]
except:
    df_final["Quarterly Score"] = 0


try:
    df_final["Yearly Score"] = df_final["Last Year negative mentions"] / df_final["Last Year total documents"]
except:
    df_final["Yearly Score"] = 0 

for i in range(len(df_final["Monthly Score"])):
    if df_final["Monthly Score"][i] > df_final["Yearly Score"][i]: 
            df_final["Urgent Need Reqd"][i] = 1
    else:
        df_final["Urgent Need Reqd"][i] = 0
    if df_final["Quarterly Score"][i] > df_final["Yearly Score"][i]:
        if df_final["Monthly Score"][i] < df_final["Quarterly Score"][i]:
            df_final["Urgent Need Reqd"][i] = 0.5
    else:
        df_final["Urgent Need Reqd"][i] = 0

df_final['Peer Tickers']=''
for i in range(len(df_final)):
    for j in range(len(df_final)):
        if i!=j:
            if df_final['Problem'][i]==df_final['Problem'][j]:
                if df_final['Revenue'][j]*0.25<df_final['Revenue'][i] or df_final['Revenue'][i]<df_final['Revenue'][j]*4:
                    df_final['Peer Tickers'][i]=df_final['Peer Tickers'][i]+' '+df_final['Ticker'][j]

df_final=df_final.drop(columns=["Last Month negative documents","Last Quarter negative documents","Last Year negative documents","Last Month total documents"])
df_final=df_final.drop(columns=["Last Quarter total documents","Last Year total documents","Last Month negative mentions","Last Quarter negative mentions"])
df_final=df_final.drop(columns=["Last Year negative mentions","Last Month total mentions","Last Quarter total mentions","Last Year total mentions"])
df_final=df_final.drop(columns=['Market Capitalization'])


document = docx.Document()
paragraph = document.add_heading('Lead Generation Tool', 0)
paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
document.add_heading('Companies selection through AlphaSense', 1)
doc_para = document.add_paragraph('This reports the companies that have been highlighted through AlphaSense as they could benefit from a project with Kearney PERLab')
document.save('doc3.docx')
doc = docx.Document('doc3.docx')
t = doc.add_table(df_final.shape[0]+1, df_final.shape[1])
t.style = 'Table Grid'
# add the header rows.
for j in range(df_final.shape[-1]):
    t.cell(0,j).text = df_final.columns[j]

# add the rest of the data frame
for i in range(df_final.shape[0]):
    for j in range(df_final.shape[-1]):
        t.cell(i+1,j).text = str(df_final.values[i,j])
# save the doc
doc.save('doc3.docx')

@app.route('/',methods=['GET'])
# @app.route('/home')
def Home():
    return render_template("index.html")





@app.route('/result',methods=['POST'])
def result():   
    # output = request.form.to_dict()
    # print(output)
    # name = output["name"]
    
    if request.method == 'POST':


        b = request.form['email']
        subject = "AlphaSense report"
        body = "Please find attached the lead generation document from AlphaSense"
        sender_email = "rr4035792@gmail.com"    #password for this email is - ramram123
        password = "ramram123"
        receiver_email = b


        # Create a multipart message and set headers
        message = MIMEMultipart()
        message["From"] = sender_email
        message["To"] = receiver_email
        message["Subject"] = subject
        message["Bcc"] = receiver_email  # Recommended for mass emails

        # Add body to email
        message.attach(MIMEText(body, "plain"))

        filename = "doc3.docx"  # In same directory as script

        # Open PDF file in binary mode
        with open(filename, "rb") as attachment:
            # Add file as application/octet-stream
            # Email client can usually download this automatically as attachment
            part = MIMEBase("application", "octet-stream")
            part.set_payload(attachment.read())

        # Encode file in ASCII characters to send by email    
        encoders.encode_base64(part)

        # Add header as key/value pair to attachment part
        part.add_header(
            "Content-Disposition",
            f"attachment; filename= {filename}",
        )

        # Add attachment to message and convert message to string
        message.attach(part)
        text = message.as_string()

        # Log in to server using secure context and send email
        context = ssl.create_default_context()
        with smtplib.SMTP_SSL("smtp.gmail.com", 465, context=context) as server:
            server.login(sender_email, password)
            server.sendmail(sender_email, receiver_email, text)


            
        output = request.form.to_dict()
        print(output)
        # name = output["name"]
        return render_template('index.html',status='Mail sent')
        


    


    return render_template('index.html')

 




if __name__ == "__main__":
    app.run(debug=True)
